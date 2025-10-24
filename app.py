import os
import json
from datetime import datetime
from functools import wraps

from flask import Flask, render_template, request, redirect, url_for, session, g, send_file, flash, abort
from flask_login import login_user, logout_user, login_required, current_user, LoginManager
import click # Importar click para comandos CLI

# Importe os formulários do seu forms.py
from forms import LoginForm, RegisterForm, EditalForm 
# Importe seus modelos de banco de dados (certifique-se de que User e Edital estão definidos em models.py)
from models import User, Edital 
# Importe as extensões (certifique-se de que extensions.py está configurado corretamente)
from extensions import db, bcrypt, login_manager, moment, csrf 

# Para manipulação de documentos .docx
from docx import Document
from docx.shared import Inches 
from docxcompose.composer import Composer 
from docx.enum.text import WD_ALIGN_PARAGRAPH 

# ================================================================
# 1. INICIALIZAÇÃO DA APLICAÇÃO FLASK
# ================================================================
app = Flask(__name__)

# Configuração da chave secreta para sessões (ESSENCIAL!)
# Altere esta chave para uma string longa e aleatória em produção
app.config['SECRET_KEY'] = 'uma_chave_secreta_muito_forte_e_aleatoria_para_producao_12345'
# Configuração do banco de dados
if os.environ.get('DATABASE_URL'):
    # Produção (Render) - PostgreSQL
    database_url = os.environ.get('DATABASE_URL')
    if database_url.startswith('postgres://'):
        database_url = database_url.replace('postgres://', 'postgresql://', 1)
    app.config['SQLALCHEMY_DATABASE_URI'] = database_url
else:
    # Desenvolvimento local - SQLite
    app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///edital_app.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False # Desativa o rastreamento de modificações para economizar recursos

# Caminhos para pastas
APP_ROOT = os.path.dirname(os.path.abspath(__file__))
GENERATED_EDITALS_FOLDER = os.path.join(APP_ROOT, 'generated_editals')
TEMPLATES_FOLDER = os.path.join(APP_ROOT, 'templates')
CLAUSULAS_FILE = os.path.join(APP_ROOT, 'clausulas.json')
MODELO_EDITAL_PATH = os.path.join(APP_ROOT, 'modelo_edital_template.docx')


# Crie a pasta generated_editals se não existir
if not os.path.exists(GENERATED_EDITALS_FOLDER):
    os.makedirs(GENERATED_EDITALS_FOLDER)

# Inicializa as extensões com o aplicativo Flask
db.init_app(app)
bcrypt.init_app(app)
login_manager.init_app(app)
moment.init_app(app)
csrf.init_app(app)

# Configuração do Flask-Login
login_manager.login_view = 'login'
login_manager.login_message_category = 'info'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# ================================================================
# 2. FUNÇÕES AUXILIARES
# ================================================================

# Decorador para exigir que o usuário seja administrador
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin():
            flash('Acesso negado. Você não tem permissões de administrador.', 'danger')
            return redirect(url_for('dashboard')) # Redireciona para o dashboard ou outra página
        return f(*args, **kwargs)
    return decorated_function

# Função para substituir placeholders em um documento DOCX
def replace_placeholder(document, placeholder, value):
    print(f"[DEBUG REPLACE] Tentando substituir '{placeholder}' por '{value}'")
    
    # Debug específico para o primeiro placeholder
    if placeholder == '{{ numero_pregao }}':
        print(f"[DEBUG FOOTER] Verificando rodapés...")
        print(f"[DEBUG FOOTER] Total de seções: {len(document.sections)}")
        for i, section in enumerate(document.sections):
            print(f"[DEBUG FOOTER] Seção {i}:")
            if section.footer:
                print(f"[DEBUG FOOTER] - Footer existe com {len(section.footer.paragraphs)} parágrafos")
                for j, para in enumerate(section.footer.paragraphs):
                    print(f"[DEBUG FOOTER] - Parágrafo {j}: '{para.text}'")
                    if '{{ numero_pregao }}' in para.text:
                        print(f"[DEBUG FOOTER] - ENCONTRADO PLACEHOLDER NO RODAPÉ!")
            else:
                print(f"[DEBUG FOOTER] - Sem footer")
    
    # Substituir nos parágrafos
    paragraphs_replaced = 0
    for paragraph in document.paragraphs:
        if placeholder in paragraph.text:
            print(f"[DEBUG REPLACE] Encontrado '{placeholder}' em parágrafo: {paragraph.text[:50]}...")
            paragraph.text = paragraph.text.replace(placeholder, str(value))
            paragraphs_replaced += 1
    
    # Substituir nas tabelas
    tables_replaced = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    print(f"[DEBUG REPLACE] Encontrado '{placeholder}' em célula: {cell.text[:50]}...")
                    cell.text = cell.text.replace(placeholder, str(value))
                    tables_replaced += 1
    
    # Substituir nos cabeçalhos e rodapés
    headers_footers_replaced = 0
    for section in document.sections:
        # Cabeçalhos
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    # Verificar texto completo do parágrafo
                    if placeholder in paragraph.text:
                        print(f"[DEBUG REPLACE] Encontrado '{placeholder}' em cabeçalho: {paragraph.text[:50]}...")
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
                        headers_footers_replaced += 1
                    else:
                        # Verificar nos runs individuais
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                print(f"[DEBUG REPLACE] Encontrado '{placeholder}' em run de cabeçalho: {run.text[:50]}...")
                                run.text = run.text.replace(placeholder, str(value))
                                headers_footers_replaced += 1
        
        # Rodapés
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    # Verificar texto completo do parágrafo
                    if placeholder in paragraph.text:
                        print(f"[DEBUG REPLACE] Encontrado '{placeholder}' em rodapé: {paragraph.text[:50]}...")
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
                        headers_footers_replaced += 1
                    else:
                        # Verificar nos runs individuais
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                print(f"[DEBUG REPLACE] Encontrado '{placeholder}' em run de rodapé: {run.text[:50]}...")
                                run.text = run.text.replace(placeholder, str(value))
                                headers_footers_replaced += 1
    
    print(f"[DEBUG REPLACE] Substituições: {paragraphs_replaced} parágrafos, {tables_replaced} células, {headers_footers_replaced} cabeçalhos/rodapés")
# ================================================================
# 3. ROTAS DA APLICAÇÃO
# ================================================================

@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user and user.check_password(form.password.data):
            login_user(user)
            flash('Login bem-sucedido!', 'success')
            next_page = request.args.get('next')
            return redirect(next_page) if next_page else redirect(url_for('dashboard'))
        else:
            flash('Usuário ou senha inválidos.', 'danger')
    return render_template('login.html', form=form)

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        # Se o usuário já está logado, redireciona para o dashboard
        flash('Você já está logado.', 'info')
        return redirect(url_for('dashboard'))
    
    form = RegisterForm()
    if form.validate_on_submit():
        # Verifica se o usuário já existe
        existing_user = User.query.filter_by(username=form.username.data).first()
        if existing_user:
            flash('Este nome de usuário já existe. Por favor, escolha outro.', 'danger')
        else:
            user = User(username=form.username.data, email=form.email.data)
            user.set_password(form.password.data)
            # A lógica de criação de admin inicial foi movida para o comando 'flask init-db'
            
            db.session.add(user)
            db.session.commit()
            flash('Registro bem-sucedido! Faça login agora.', 'success')
            return redirect(url_for('login'))
    return render_template('register.html', form=form)

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Você foi desconectado.', 'info')
    return redirect(url_for('login'))

@app.route('/dashboard')
@login_required
def dashboard():
    # Carrega os editais do usuário logado, ordenados pela data de criação
    editals = Edital.query.filter_by(creator_id=current_user.id).order_by(Edital.data_criacao.desc()).all()
    return render_template('dashboard.html', editals=editals)
def processar_tipo_participacao(tipo_participacao):
    if tipo_participacao == 'ampla':
        return '''1( X ) LICITAÇÃO DE AMPLA PARTICIPAÇÃO.
O item 2.1 alínea "b" das Condições Específicas do Edital não é aplicável.

2(   ) LICITAÇÃO DE PARTICIPAÇÃO EXCLUSIVA DE MICROEMPRESAS, EMPRESAS DE PEQUENO PORTE OU COOPERATIVAS QUE PREENCHAM AS CONDIÇÕES ESTABELECIDAS NO ARTIGO 34 DA LEI FEDERAL Nº 11.488, DE 15/06/2007.'''
    
    elif tipo_participacao == 'exclusiva':
        return '''1(   ) LICITAÇÃO DE AMPLA PARTICIPAÇÃO.
O item 2.1 alínea "b" das Condições Específicas do Edital não é aplicável.

2( X ) LICITAÇÃO DE PARTICIPAÇÃO EXCLUSIVA DE MICROEMPRESAS, EMPRESAS DE PEQUENO PORTE OU COOPERATIVAS QUE PREENCHAM AS CONDIÇÕES ESTABELECIDAS NO ARTIGO 34 DA LEI FEDERAL Nº 11.488, DE 15/06/2007.'''
    
    return ''

@app.route('/generate_edital', methods=['GET', 'POST'])
@login_required
def generate_edital():
    form = EditalForm() # Instancia o formulário WTForms
    clausulas_data = {}
    try:
        with open(CLAUSULAS_FILE, 'r', encoding='utf-8') as f:
            clausulas_data = json.load(f)
    except FileNotFoundError:
        flash(f'Arquivo de cláusulas não encontrado: {CLAUSULAS_FILE}', 'danger')
    except json.JSONDecodeError:
        flash(f'Erro ao ler o arquivo JSON de cláusulas: {CLAUSULAS_FILE}', 'danger')

    if form.validate_on_submit(): # Usa validate_on_submit para WTForms
        try:
            # Coleta de dados do formulário usando form.field.data
            form_name = form.form_name.data
            numero_pregao = form.numero_pregao.data
            objeto_servicos = form.objeto_servicos.data
            tipo_participacao = form.tipo_participacao.data
            clausula_participacao = processar_tipo_participacao(tipo_participacao)
            
            # Novos campos
            compras_gov_numero = form.compras_gov_numero.data
            valor_total_orcamento = form.valor_total_orcamento.data
            data_base_orcamento = form.data_base_orcamento.data
            data_sessao = form.data_sessao.data
            hora_sessao = form.hora_sessao.data
            data_disponibilidade = form.data_disponibilidade.data
            email_contato1 = form.email_contato1.data
            email_contato2 = form.email_contato2.data
            orcamento_sigiloso = form.orcamento_sigiloso.data

            # Campos do formulário que serão persistidos
            permite_visita_tecnica = form.permite_visita_tecnica.data
            criterio_julgamento = form.criterio_julgamento.data
            aplicacao_criterio = form.aplicacao_criterio.data
            modo_disputa = form.modo_disputa.data
            tipo_participacao = form.tipo_participacao.data
            participacao_consorcio = form.participacao_consorcio.data
            diferencial_aliquota = form.diferencial_aliquota.data 
            regularidade_fiscal = form.regularidade_fiscal.data
            qualificacao_tecnica = form.qualificacao_tecnica.data
            atestados_qualificacao_tecnica = form.atestados_qualificacao_tecnica.data
            qualificacao_economico_financeira = form.qualificacao_economico_financeira.data
            servico_continuo = form.servico_continuo.data
            garantia_sim_nao = form.garantia_sim_nao.data
            subcontratacao = form.subcontratacao.data
            permitido_cooperativa = form.permitido_cooperativa.data
            cad_madeira = form.cad_madeira.data
            documento_tecnico_sim_nao = form.documento_tecnico_sim_nao.data
            documento_tecnico_nome = form.documento_tecnico_nome.data

            # Campos do Anexo 1
            numero_licitacao_anexo1 = form.numero_licitacao_anexo1.data
            objeto_licitacao_anexo1 = form.objeto_licitacao_anexo1.data

            # Campos booleanos para as declarações do Anexo 1
            incluir_rec_judicial = form.incluir_rec_judicial.data
            incluir_rec_extrajudicial = form.incluir_rec_extrajudicial.data
            incluir_me_epp = form.incluir_me_epp.data
            incluir_cadmadeira = form.incluir_cadmadeira.data

            # Campos de Cláusulas Específicas
            regime_empreitada = form.regime_empreitada.data
            prazos_execucao = form.prazos_execucao.data
            tipo_instrumento_contratual = form.tipo_instrumento_contratual.data
            prorrogacao_contrato = form.prorrogacao_contrato.data
            medicao_servicos = form.medicao_servicos.data
            fiscalizacao_inspecao = form.fiscalizacao_inspecao.data
            consequencias_rescisao = form.consequencias_rescisao.data
            suspensao_temporaria_servicos = form.suspensao_temporaria_servicos.data
            aceitacao_servicos = form.aceitacao_servicos.data
            garantia_servicos = form.garantia_servicos.data


            # Cria uma nova instância de Edital
            new_edital = Edital(
                form_name=form_name,
                numero_pregao=numero_pregao,
                objeto_servicos=objeto_servicos,
                compras_gov_numero=compras_gov_numero,
                valor_total_orcamento=valor_total_orcamento,
                data_base_orcamento=data_base_orcamento,
                data_sessao=data_sessao,
                hora_sessao=hora_sessao,
                data_disponibilidade=data_disponibilidade,
                email_contato1=email_contato1,
                email_contato2=email_contato2,
                orcamento_sigiloso=orcamento_sigiloso,
                creator_id=current_user.id, # Associa o edital ao usuário logado
                permite_visita_tecnica=permite_visita_tecnica,
                criterio_julgamento=criterio_julgamento,
                aplicacao_criterio=aplicacao_criterio,
                modo_disputa=modo_disputa,
                tipo_participacao=tipo_participacao,
                participacao_consorcio=participacao_consorcio,
                diferencial_aliquota=diferencial_aliquota, 
                regularidade_fiscal=regularidade_fiscal,
                qualificacao_tecnica=qualificacao_tecnica,
                atestados_qualificacao_tecnica=atestados_qualificacao_tecnica,
                qualificacao_economico_financeira=qualificacao_economico_financeira,
                servico_continuo=servico_continuo,
                garantia_sim_nao=garantia_sim_nao,
                subcontratacao=subcontratacao,
                permitido_cooperativa=permitido_cooperativa,
                cad_madeira=cad_madeira,
                documento_tecnico_sim_nao=documento_tecnico_sim_nao,
                documento_tecnico_nome=documento_tecnico_nome,
                numero_licitacao_anexo1=numero_licitacao_anexo1,
                objeto_licitacao_anexo1=objeto_licitacao_anexo1,
                incluir_rec_judicial=incluir_rec_judicial,
                incluir_rec_extrajudicial=incluir_rec_extrajudicial,
                incluir_me_epp=incluir_me_epp,
                incluir_cadmadeira=incluir_cadmadeira,
                regime_empreitada=regime_empreitada,
                prazos_execucao=prazos_execucao,
                tipo_instrumento_contratual=tipo_instrumento_contratual,
                prorrogacao_contrato=prorrogacao_contrato,
                medicao_servicos=medicao_servicos,
                fiscalizacao_inspecao=fiscalizacao_inspecao,
                consequencias_rescisao=consequencias_rescisao,
                suspensao_temporaria_servicos=suspensao_temporaria_servicos,
                aceitacao_servicos=aceitacao_servicos,
                garantia_servicos=garantia_servicos
            )
            
            db.session.add(new_edital)
            db.session.commit()

            # Lógica para preencher o template .docx
            print(f"[DEBUG] Caminho do modelo DOCX: {MODELO_EDITAL_PATH}")
            if not os.path.exists(MODELO_EDITAL_PATH):
                raise FileNotFoundError(f"Modelo de edital não encontrado em: {MODELO_EDITAL_PATH}")

            document = Document(MODELO_EDITAL_PATH)
            
            # Mapeamento dos campos do formulário para os placeholders no DOCX
            # Use o dicionário para facilitar a substituição
            replacements = {
                '{{ numero_pregao }}': numero_pregao,
                '{{ objeto_servicos }}': objeto_servicos,
                '{{ compras_gov_numero }}': str(compras_gov_numero) if compras_gov_numero else '',
                '{{ valor_total_contratacao }}': str(valor_total_orcamento) if valor_total_orcamento else '',
                '{{ critério_julgamento_resumo }}': clausulas_data['criterio_julgamento'].get(criterio_julgamento, '').upper(),
                '{{ item_grupo_global_resumo }}': clausulas_data['aplicacao_criterio'].get(aplicacao_criterio, '').upper(),
                '{{ modo_disputa_resumo }}': clausulas_data['modo_disputa'].get(modo_disputa, '').upper(),
                '{{ data_sessao }}': data_sessao.strftime('%d/%m/%Y') if data_sessao else '',
                '{{ hora_sessao }}': hora_sessao if hora_sessao else '',
                '{{ data_disponibilidade }}': data_disponibilidade.strftime('%d/%m/%Y') if data_disponibilidade else '',
                '{{ documento_tecnico }}': documento_tecnico_nome if documento_tecnico_sim_nao == 'sim' and documento_tecnico_nome else '(QUANDO COUBER)',
                '{{ licitação_ampla }}': 'X' if tipo_participacao == 'ampla' else '',
                '{{ clausula_participacao }}': clausula_participacao,  # ← ADICIONE ESTA LINHA
                '{{ licitação_micro }}': 'X' if tipo_participacao == 'micro' else '',
                '{{ numero_licitacao_anexo1 }}': str(numero_licitacao_anexo1) if numero_licitacao_anexo1 else '',
                '{{ objeto_licitacao_anexo1 }}': objeto_licitacao_anexo1,
                '{{ declaração_rec_judicial }}': clausulas_data['declaracoes_anexo1']['recuperacao_judicial'] if incluir_rec_judicial else '',
                '{{ declaração_rec_extrajudicial }}': clausulas_data['declaracoes_anexo1']['recuperacao_extrajudicial'] if incluir_rec_extrajudicial else '',
                '{{ declaração_me_epp }}': clausulas_data['declaracoes_anexo1']['micro_empresa_epp'] if incluir_me_epp else '',
                '{{ declaração_cadmadeira }}': clausulas_data['declaracoes_anexo1']['cadmadeira'] if incluir_cadmadeira else '',
                '{{ maior_desconto }}': clausulas_data['criterio_julgamento']['maior_desconto'] if criterio_julgamento == 'maior_desconto' else '',
                '{{ menor_preço }}': clausulas_data['criterio_julgamento']['menor_preco'] if criterio_julgamento == 'menor_preco' else '',
                '{{ participação_cooperativas }}': clausulas_data['permitido_cooperativa'].get(permitido_cooperativa, ''),
                '{{ participação_consorcio }}': clausulas_data['participacao_consorcio'].get(participacao_consorcio, ''),
                '{{ não_participação_consorcio }}': clausulas_data['nao_participacao_consorcio'] if participacao_consorcio == 'nao' else '',
                '{{ proposta_maior_desconto }}': clausulas_data['proposta_maior_desconto'] if criterio_julgamento == 'maior_desconto' else '',
                '{{ com_material }}': clausulas_data['diferencial_aliquota'].get(diferencial_aliquota, ''), 
                '{{ prova_regularidade_fical }}': clausulas_data['regularidade_fiscal'].get(regularidade_fiscal, ''),
                '{{ qualificação_tecnica }}': clausulas_data['qualificacao_tecnica'].get(qualificacao_tecnica, ''),
                '{{ exigência_prazo }}': clausulas_data['atestados_qualificacao_tecnica'].get(atestados_qualificacao_tecnica, ''),
                '{{ visita_tecnica }}': clausulas_data['permite_visita_tecnica'].get(permite_visita_tecnica, ''),
                '{{ certidão_negativa }}': clausulas_data['qualificacao_economico_financeira']['exigir']['certidao_negativa'] if qualificacao_economico_financeira == 'exigir' else '',
                '{{ balanço_patrimonial }}': clausulas_data['qualificacao_economico_financeira']['exigir']['balanco_patrimonial'] if qualificacao_economico_financeira == 'exigir' else clausulas_data['qualificacao_economico_financeira']['nao_exigir']['balanco_patrimonial'],
                '{{ índice_liquidez }}': clausulas_data['qualificacao_economico_financeira']['exigir']['indice_liquidez'] if qualificacao_economico_financeira == 'exigir' else '',
                '{{ patrimônio_liquido }}': clausulas_data['qualificacao_economico_financeira']['exigir']['patrimonio_liquido'] if qualificacao_economico_financeira == 'exigir' else '',
                '{{ valor_percentual }}': clausulas_data['valor_percentual'],
                '{{ aberto_fechado_ambos }}': clausulas_data['modo_disputa'].get(modo_disputa, '').upper(),
                '{{ maior_menor_pregao }}': clausulas_data['julgamento_pregao'].get(criterio_julgamento, ''),
                '{{ oferta_julgamento_resumo }}': clausulas_data['oferta_julgamento_resumo'].get(criterio_julgamento, ''),
                '{{ menor_maior_oferta }}': clausulas_data['menor_maior_oferta'].get(criterio_julgamento, ''),
                '{{ maior_desconto_escolha }}': clausulas_data['contratacao_escolha']['maior_desconto'] if criterio_julgamento == 'maior_desconto' else '',
                '{{ menor-preço_escolha }}': clausulas_data['contratacao_escolha']['menor_preco'] if criterio_julgamento == 'menor_preco' else '',
                '{{ garantia_execução }}': clausulas_data['garantia_execucao'].get(garantia_sim_nao, ''),
                '{{ sub_contratação }}': clausulas_data['subcontratacao'].get(subcontratacao, ''),
                '{{ cooperativa_gestor }}': clausulas_data['permitido_cooperativa'].get(permitido_cooperativa, '') if permitido_cooperativa == 'sim' else '', 
                '{{ certidão_negativa_administrador }}': clausulas_data['certidao_negativa_administrador'] if qualificacao_economico_financeira == 'exigir' else '',
                '{{ madeira }}': clausulas_data['cad_madeira_detalhe'] if cad_madeira == 'sim' else '',
                '{{ fiscalização_inspecao }}': clausulas_data['fiscalizacao_inspecao_contrato'].get(fiscalizacao_inspecao, ''),
                '{{ orçamento_sigiloso }}': clausulas_data['orcamento_sigiloso_texto'] if orcamento_sigiloso == 'sim' else '',
                '{{ edital_condicionais.isento_icms_completa }}': '', # Mantido vazio conforme discutido
                '{{ email_contato1 }}': email_contato1 if email_contato1 else 'email1@exemplo.com',
                '{{ email_contato2 }}': email_contato2 if email_contato2 else 'email2@exemplo.com',
                '{{ nome }}': current_user.username,
                '{{ cargo }}': 'Gerente de Projetos', # Exemplo de cargo, pode ser dinâmico
                '{{ instrumento_contratual }}': clausulas_data['tipo_instrumento_contratual_contrato'].get(tipo_instrumento_contratual, ''),
                '{{ regime_empreitada }}': clausulas_data['regime_empreitada_contrato'].get(regime_empreitada, ''),
                '{{ prazos_execucao }}': clausulas_data['prazos_execucao_contrato'].get(prazos_execucao, ''),
                '{{ tipo_instrumento_contratual }}': clausulas_data['tipo_instrumento_contratual_contrato'].get(tipo_instrumento_contratual, ''),
                '{{ prorrogacao_contrato }}': clausulas_data['prorrogacao_contrato_contrato'].get(prorrogacao_contrato, ''),
                '{{ medicao_servicos }}': clausulas_data['medicao_servicos_contrato'].get(medicao_servicos, ''),
                '{{ fiscalizacao_inspecao }}': clausulas_data['fiscalizacao_inspecao_contrato'].get(fiscalizacao_inspecao, ''),
                '{{ consequencias_rescisao }}': clausulas_data['consequencias_rescisao_contrato'].get(consequencias_rescisao, ''),
                '{{ suspensao_temporaria_servicos }}': clausulas_data['suspensao_temporaria_servicos_contrato'].get(suspensao_temporaria_servicos, ''),
                '{{ aceitacao_servicos }}': clausulas_data['aceitacao_servicos_contrato'].get(aceitacao_servicos, ''),
                '{{ garantia_servicos }}': clausulas_data['garantia_servicos_contrato'].get(garantia_servicos, ''),
                '{{ nome_usuario }}': current_user.username, # Adicionado para o TCN
                '{{ cargo_usuario }}': 'Analista', # Adicionado para o TCN, pode ser dinâmico
            }

            print(f"[DEBUG] Iniciando substituição de placeholders...")
            print(f"[DEBUG] Total de placeholders: {len(replacements)}")
            for placeholder, value in replacements.items():
                print(f"[DEBUG] Substituindo '{placeholder}' por '{value}'")
                replace_placeholder(document, placeholder, value)
            print(f"[DEBUG] Substituição concluída!")

            # Salve o novo documento
            filename = f"Edital_{form_name.replace(' ', '_').replace('/', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(GENERATED_EDITALS_FOLDER, filename)

            # Salve o novo documento
            filename = f"Edital_{form_name.replace(' ', '_').replace('/', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(GENERATED_EDITALS_FOLDER, filename)
            
            print(f"[DEBUG] Nome do arquivo gerado: {filename}")
            print(f"[DEBUG] Caminho completo do arquivo: {filepath}")

            document.save(filepath)
            print(f"[DEBUG SAVE] Documento salvo com {len(document.paragraphs)} parágrafos")
            print(f"[DEBUG SAVE] Verificando se '12345' está no documento salvo...")
            for i, para in enumerate(document.paragraphs[:10]):
             if '12345' in para.text:
              print(f"[DEBUG SAVE] Encontrado '12345' no parágrafo {i}: {para.text[:100]}")

            # Atualiza o caminho do arquivo gerado no banco de dados
            new_edital.generated_filename = filename
            db.session.commit()
            print(f"[DEBUG] Nome do arquivo '{filename}' salvo no banco de dados para o edital ID: {new_edital.id}")


            flash(f'Edital "{form_name}" gerado com sucesso!', 'success')
            return redirect(url_for('dashboard'))

        except FileNotFoundError:
            flash(f'Arquivo de template não encontrado: {MODELO_EDITAL_PATH}', 'danger')
            print(f"[ERROR] FileNotFoundError: {MODELO_EDITAL_PATH}")
        except Exception as e:
            flash(f'Erro ao gerar o edital: {str(e)}', 'danger')
            # Opcional: logar o erro completo para depuração
            app.logger.error(f"Erro na geração do edital: {e}", exc_info=True)
            print(f"[ERROR] Erro inesperado durante a geração do edital: {e}")

    return render_template('generate_edital.html', form=form, clausulas=clausulas_data)

# Rota para edição de edital
@app.route('/edit_edital/<int:edital_id>', methods=['GET', 'POST'])
@login_required
def edit_edital(edital_id):
    edital = Edital.query.get_or_404(edital_id)

    # Autorização
        # Autorização: Apenas o criador OU um admin pode editar
    if edital.creator_id != current_user.id and not current_user.is_admin():
        flash('Você não tem permissão para editar este edital.', 'danger')
        return redirect(url_for('dashboard'))

    form = EditalForm()
    clausulas_data = {}
    try:
        with open(CLAUSULAS_FILE, 'r', encoding='utf-8') as f:
            clausulas_data = json.load(f)
    except FileNotFoundError:
        flash(f'Arquivo de cláusulas não encontrado: {CLAUSULAS_FILE}', 'danger')
    except json.JSONDecodeError:
        flash(f'Erro ao ler o arquivo JSON de cláusulas: {CLAUSULAS_FILE}', 'danger')

    if form.validate_on_submit():
        try:
            # Armazena o nome do arquivo antigo para possível exclusão
            old_filename = edital.generated_filename

            # Atualiza o objeto edital com os dados do formulário
            edital.form_name = form.form_name.data
            edital.numero_pregao = form.numero_pregao.data
            edital.objeto_servicos = form.objeto_servicos.data
            edital.compras_gov_numero = form.compras_gov_numero.data
            edital.valor_total_orcamento = form.valor_total_orcamento.data
            edital.data_base_orcamento = form.data_base_orcamento.data
            edital.data_sessao = form.data_sessao.data
            edital.hora_sessao = form.hora_sessao.data
            edital.data_disponibilidade = form.data_disponibilidade.data
            edital.email_contato1 = form.email_contato1.data
            edital.email_contato2 = form.email_contato2.data
            edital.orcamento_sigiloso = form.orcamento_sigiloso.data
            edital.permite_visita_tecnica = form.permite_visita_tecnica.data
            edital.criterio_julgamento = form.criterio_julgamento.data
            edital.aplicacao_criterio = form.aplicacao_criterio.data
            edital.modo_disputa = form.modo_disputa.data
            edital.tipo_participacao = form.tipo_participacao.data
            edital.participacao_consorcio = form.participacao_consorcio.data
            edital.diferencial_aliquota = form.diferencial_aliquota.data
            edital.regularidade_fiscal = form.regularidade_fiscal.data
            edital.qualificacao_tecnica = form.qualificacao_tecnica.data
            edital.atestados_qualificacao_tecnica = form.atestados_qualificacao_tecnica.data
            edital.qualificacao_economico_financeira = form.qualificacao_economico_financeira.data
            edital.servico_continuo = form.servico_continuo.data
            edital.garantia_sim_nao = form.garantia_sim_nao.data
            edital.subcontratacao = form.subcontratacao.data
            edital.permitido_cooperativa = form.permitido_cooperativa.data
            edital.cad_madeira = form.cad_madeira.data
            edital.documento_tecnico_sim_nao = form.documento_tecnico_sim_nao.data
            edital.documento_tecnico_nome = form.documento_tecnico_nome.data
            edital.numero_licitacao_anexo1 = form.numero_licitacao_anexo1.data
            edital.objeto_licitacao_anexo1 = form.objeto_licitacao_anexo1.data
            edital.incluir_rec_judicial = form.incluir_rec_judicial.data
            edital.incluir_rec_extrajudicial = form.incluir_rec_extrajudicial.data
            edital.incluir_me_epp = form.incluir_me_epp.data
            edital.incluir_cadmadeira = form.incluir_cadmadeira.data
            edital.regime_empreitada = form.regime_empreitada.data
            edital.prazos_execucao = form.prazos_execucao.data
            edital.tipo_instrumento_contratual = form.tipo_instrumento_contratual.data
            edital.prorrogacao_contrato = form.prorrogacao_contrato.data
            edital.medicao_servicos = form.medicao_servicos.data
            edital.fiscalizacao_inspecao = form.fiscalizacao_inspecao.data
            edital.consequencias_rescisao = form.consequencias_rescisao.data
            edital.suspensao_temporaria_servicos = form.suspensao_temporaria_servicos.data
            edital.aceitacao_servicos = form.aceitacao_servicos.data
            edital.garantia_servicos = form.garantia_servicos.data

            # --- INÍCIO DA LÓGICA DE RE-GERAÇÃO DO DOCX (Copiado de generate_edital) ---
            document = Document(MODELO_EDITAL_PATH)
            
            replacements = {
                '{{ numero_pregao }}': edital.numero_pregao,
                '{{ objeto_servicos }}': edital.objeto_servicos,
                '{{ compras_gov_numero }}': str(edital.compras_gov_numero) if edital.compras_gov_numero else '',
                '{{ valor_total_contratacao }}': str(edital.valor_total_orcamento) if edital.valor_total_orcamento else '',
                '{{ critério_julgamento_resumo }}': clausulas_data['criterio_julgamento'].get(edital.criterio_julgamento, '').upper(),
                '{{ item_grupo_global_resumo }}': clausulas_data['aplicacao_criterio'].get(edital.aplicacao_criterio, '').upper(),
                '{{ modo_disputa_resumo }}': clausulas_data['modo_disputa'].get(edital.modo_disputa, '').upper(),
                '{{ data_sessao }}': edital.data_sessao.strftime('%d/%m/%Y') if edital.data_sessao else '',
                '{{ hora_sessao }}': edital.hora_sessao if edital.hora_sessao else '',
                '{{ clausula_participacao }}': processar_tipo_participacao(edital.tipo_participacao),
                '{{ data_disponibilidade }}': edital.data_disponibilidade.strftime('%d/%m/%Y') if edital.data_disponibilidade else '',
                '{{ documento_tecnico }}': edital.documento_tecnico_nome if edital.documento_tecnico_sim_nao == 'sim' and edital.documento_tecnico_nome else '(QUANDO COUBER)',
                '{{ licitação_ampla }}': 'X' if edital.tipo_participacao == 'ampla' else '',
                '{{ licitação_micro }}': 'X' if edital.tipo_participacao == 'micro' else '',
                '{{ numero_licitacao_anexo1 }}': str(edital.numero_licitacao_anexo1) if edital.numero_licitacao_anexo1 else '',
                '{{ objeto_licitacao_anexo1 }}': edital.objeto_licitacao_anexo1,
                '{{ declaração_rec_judicial }}': clausulas_data['declaracoes_anexo1']['recuperacao_judicial'] if edital.incluir_rec_judicial else '',
                '{{ declaração_rec_extrajudicial }}': clausulas_data['declaracoes_anexo1']['recuperacao_extrajudicial'] if edital.incluir_rec_extrajudicial else '',
                '{{ declaração_me_epp }}': clausulas_data['declaracoes_anexo1']['micro_empresa_epp'] if edital.incluir_me_epp else '',
                '{{ declaração_cadmadeira }}': clausulas_data['declaracoes_anexo1']['cadmadeira'] if edital.incluir_cadmadeira else '',
                '{{ maior_desconto }}': clausulas_data['criterio_julgamento']['maior_desconto'] if edital.criterio_julgamento == 'maior_desconto' else '',
                '{{ menor_preço }}': clausulas_data['criterio_julgamento']['menor_preco'] if edital.criterio_julgamento == 'menor_preco' else '',
                '{{ participação_cooperativas }}': clausulas_data['permitido_cooperativa'].get(edital.permitido_cooperativa, ''),
                '{{ participação_consorcio }}': clausulas_data['participacao_consorcio'].get(edital.participacao_consorcio, ''),
                '{{ não_participação_consorcio }}': clausulas_data['nao_participacao_consorcio'] if edital.participacao_consorcio == 'nao' else '',
                '{{ proposta_maior_desconto }}': clausulas_data['proposta_maior_desconto'] if edital.criterio_julgamento == 'maior_desconto' else '',
                '{{ com_material }}': clausulas_data['diferencial_aliquota'].get(edital.diferencial_aliquota, ''), 
                '{{ prova_regularidade_fical }}': clausulas_data['regularidade_fiscal'].get(edital.regularidade_fiscal, ''),
                '{{ qualificação_tecnica }}': clausulas_data['qualificacao_tecnica'].get(edital.qualificacao_tecnica, ''),
                '{{ exigência_prazo }}': clausulas_data['atestados_qualificacao_tecnica'].get(edital.atestados_qualificacao_tecnica, ''),
                '{{ visita_tecnica }}': clausulas_data['permite_visita_tecnica'].get(edital.permite_visita_tecnica, ''),
                '{{ certidão_negativa }}': clausulas_data['qualificacao_economico_financeira']['exigir']['certidao_negativa'] if edital.qualificacao_economico_financeira == 'exigir' else '',
                '{{ balanço_patrimonial }}': clausulas_data['qualificacao_economico_financeira']['exigir']['balanco_patrimonial'] if edital.qualificacao_economico_financeira == 'exigir' else clausulas_data['qualificacao_economico_financeira']['nao_exigir']['balanco_patrimonial'],
                '{{ índice_liquidez }}': clausulas_data['qualificacao_economico_financeira']['exigir']['indice_liquidez'] if edital.qualificacao_economico_financeira == 'exigir' else '',
                '{{ patrimônio_liquido }}': clausulas_data['qualificacao_economico_financeira']['exigir']['patrimonio_liquido'] if edital.qualificacao_economico_financeira == 'exigir' else '',
                '{{ valor_percentual }}': clausulas_data['valor_percentual'],
                '{{ aberto_fechado_ambos }}': clausulas_data['modo_disputa'].get(edital.modo_disputa, '').upper(),
                '{{ maior_menor_pregao }}': clausulas_data['julgamento_pregao'].get(edital.criterio_julgamento, ''),
                '{{ oferta_julgamento_resumo }}': clausulas_data['oferta_julgamento_resumo'].get(edital.criterio_julgamento, ''),
                '{{ menor_maior_oferta }}': clausulas_data['menor_maior_oferta'].get(edital.criterio_julgamento, ''),
                '{{ maior_desconto_escolha }}': clausulas_data['contratacao_escolha']['maior_desconto'] if edital.criterio_julgamento == 'maior_desconto' else '',
                '{{ menor-preço_escolha }}': clausulas_data['contratacao_escolha']['menor_preco'] if edital.criterio_julgamento == 'menor_preco' else '',
                '{{ garantia_execução }}': clausulas_data['garantia_execucao'].get(edital.garantia_sim_nao, ''),
                '{{ sub_contratação }}': clausulas_data['subcontratacao'].get(edital.subcontratacao, ''),
                '{{ cooperativa_gestor }}': clausulas_data['permitido_cooperativa'].get(edital.permitido_cooperativa, '') if edital.permitido_cooperativa == 'sim' else '', 
                '{{ certidão_negativa_administrador }}': clausulas_data['certidao_negativa_administrador'] if edital.qualificacao_economico_financeira == 'exigir' else '',
                '{{ madeira }}': clausulas_data['cad_madeira_detalhe'] if edital.cad_madeira == 'sim' else '',
                '{{ fiscalização_inspecao }}': clausulas_data['fiscalizacao_inspecao_contrato'].get(edital.fiscalizacao_inspecao, ''),
                '{{ orçamento_sigiloso }}': clausulas_data['orcamento_sigiloso_texto'] if edital.orcamento_sigiloso == 'sim' else '',
                '{{ edital_condicionais.isento_icms_completa }}': '',
                '{{ email_contato1 }}': edital.email_contato1 if edital.email_contato1 else 'email1@exemplo.com',
                '{{ email_contato2 }}': edital.email_contato2 if edital.email_contato2 else 'email2@exemplo.com',
                '{{ nome }}': current_user.username,
                '{{ cargo }}': 'Gerente de Projetos',
                '{{ instrumento_contratual }}': clausulas_data['tipo_instrumento_contratual_contrato'].get(edital.tipo_instrumento_contratual, ''),
                '{{ regime_empreitada }}': clausulas_data['regime_empreitada_contrato'].get(edital.regime_empreitada, ''),
                '{{ prazos_execucao }}': clausulas_data['prazos_execucao_contrato'].get(edital.prazos_execucao, ''),
                '{{ tipo_instrumento_contratual }}': clausulas_data['tipo_instrumento_contratual_contrato'].get(edital.tipo_instrumento_contratual, ''),
                '{{ prorrogacao_contrato }}': clausulas_data['prorrogacao_contrato_contrato'].get(edital.prorrogacao_contrato, ''),
                '{{ medicao_servicos }}': clausulas_data['medicao_servicos_contrato'].get(edital.medicao_servicos, ''),
                '{{ fiscalizacao_inspecao }}': clausulas_data['fiscalizacao_inspecao_contrato'].get(edital.fiscalizacao_inspecao, ''),
                '{{ consequencias_rescisao }}': clausulas_data['consequencias_rescisao_contrato'].get(edital.consequencias_rescisao, ''),
                '{{ suspensao_temporaria_servicos }}': clausulas_data['suspensao_temporaria_servicos_contrato'].get(edital.suspensao_temporaria_servicos, ''),
                '{{ aceitacao_servicos }}': clausulas_data['aceitacao_servicos_contrato'].get(edital.aceitacao_servicos, ''),
                '{{ garantia_servicos }}': clausulas_data['garantia_servicos_contrato'].get(edital.garantia_servicos, ''),
                '{{ nome_usuario }}': current_user.username,
                '{{ cargo_usuario }}': 'Analista',
            }

            print(f"[DEBUG] Iniciando substituição de placeholders...")
            for placeholder, value in replacements.items():
             replace_placeholder(document, placeholder, value)

            # Gera um novo nome de arquivo para o edital editado
            new_generated_filename = f"Edital_{edital.form_name.replace(' ', '_').replace('/', '_')}_EDITED_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            new_filepath = os.path.join(GENERATED_EDITALS_FOLDER, new_generated_filename)
            document.save(new_filepath)

            # Atualiza o objeto edital com o novo nome de arquivo gerado
            edital.generated_filename = new_generated_filename
            db.session.commit() # Confirma as alterações no banco de dados, incluindo o novo nome do arquivo

            # Opcional: Exclui o arquivo antigo se ele existia e é diferente do novo
            if old_filename and old_filename != new_generated_filename:
                old_filepath = os.path.join(GENERATED_EDITALS_FOLDER, old_filename)
                if os.path.exists(old_filepath):
                    try:
                        os.remove(old_filepath)
                        flash(f'Arquivo antigo {old_filename} removido.', 'info')
                    except Exception as e:
                        flash(f'Erro ao remover arquivo antigo {old_filepath}: {str(e)}', 'warning')
                        app.logger.error(f"Erro ao remover arquivo antigo {old_filepath}: {e}", exc_info=True)
            # --- FIM DA LÓGICA DE RE-GERAÇÃO DO DOCX ---

            flash('Edital atualizado e arquivo DOCX re-gerado com sucesso!', 'success')
            return redirect(url_for('dashboard'))

        except FileNotFoundError:
            flash(f'Arquivo de template não encontrado: {MODELO_EDITAL_PATH}', 'danger')
        except Exception as e:
            flash(f'Erro ao atualizar o edital e re-gerar o arquivo: {str(e)}', 'danger')
            app.logger.error(f"Erro na atualização e re-geração do edital: {e}", exc_info=True)

    elif request.method == 'GET':
        # Pré-preenche os campos do formulário com os dados do edital existente
        form.form_name.data = edital.form_name
        form.numero_pregao.data = edital.numero_pregao
        form.objeto_servicos.data = edital.objeto_servicos
        form.compras_gov_numero.data = edital.compras_gov_numero
        form.valor_total_orcamento.data = edital.valor_total_orcamento
        form.data_base_orcamento.data = edital.data_base_orcamento
        form.data_sessao.data = edital.data_sessao
        form.hora_sessao.data = edital.hora_sessao
        form.data_disponibilidade.data = edital.data_disponibilidade
        form.email_contato1.data = edital.email_contato1
        form.email_contato2.data = edital.email_contato2
        form.orcamento_sigiloso.data = edital.orcamento_sigiloso
        form.permite_visita_tecnica.data = edital.permite_visita_tecnica
        form.criterio_julgamento.data = edital.criterio_julgamento
        form.aplicacao_criterio.data = edital.aplicacao_criterio
        form.modo_disputa.data = edital.modo_disputa
        form.tipo_participacao.data = edital.tipo_participacao
        form.participacao_consorcio.data = edital.participacao_consorcio
        form.diferencial_aliquota.data = edital.diferencial_aliquota
        form.regularidade_fiscal.data = edital.regularidade_fiscal
        form.qualificacao_tecnica.data = edital.qualificacao_tecnica
        form.atestados_qualificacao_tecnica.data = edital.atestados_qualificacao_tecnica
        form.qualificacao_economico_financeira.data = edital.qualificacao_economico_financeira
        form.servico_continuo.data = edital.servico_continuo
        form.garantia_sim_nao.data = edital.garantia_sim_nao
        form.subcontratacao.data = edital.subcontratacao
        form.permitido_cooperativa.data = edital.permitido_cooperativa
        form.cad_madeira.data = edital.cad_madeira
        form.documento_tecnico_sim_nao.data = edital.documento_tecnico_sim_nao
        form.documento_tecnico_nome.data = edital.documento_tecnico_nome
        form.numero_licitacao_anexo1.data = edital.numero_licitacao_anexo1
        form.objeto_licitacao_anexo1.data = edital.objeto_licitacao_anexo1
        form.incluir_rec_judicial.data = edital.incluir_rec_judicial
        form.incluir_rec_extrajudicial.data = edital.incluir_rec_extrajudicial
        form.incluir_me_epp.data = edital.incluir_me_epp
        form.incluir_cadmadeira.data = edital.incluir_cadmadeira
        form.regime_empreitada.data = edital.regime_empreitada
        form.prazos_execucao.data = edital.prazos_execucao
        form.tipo_instrumento_contratual.data = edital.tipo_instrumento_contratual
        form.prorrogacao_contrato.data = edital.prorrogacao_contrato
        form.medicao_servicos.data = edital.medicao_servicos
        form.fiscalizacao_inspecao.data = edital.fiscalizacao_inspecao
        form.consequencias_rescisao.data = edital.consequencias_rescisao
        form.suspensao_temporaria_servicos.data = edital.suspensao_temporaria_servicos
        form.aceitacao_servicos.data = edital.aceitacao_servicos
        form.garantia_servicos.data = edital.garantia_servicos

    return render_template('generate_edital.html', form=form, clausulas=clausulas_data, edital_id=edital_id)


@app.route('/download_edital/<filename>')
@login_required
def download_edital(filename):
    # Permite que o criador OU um admin baixe o edital
    # Primeiro, encontra o edital pelo filename para verificar o criador
    edital = Edital.query.filter_by(generated_filename=filename).first()
    if not edital:
        flash('Edital não encontrado no banco de dados.', 'danger')
        return redirect(url_for('dashboard'))

    if edital.creator_id != current_user.id and not current_user.is_admin():
        flash('Você não tem permissão para baixar este edital.', 'danger')
        return redirect(url_for('dashboard'))

    filepath = os.path.join(GENERATED_EDITALS_FOLDER, filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    else:
        flash('Arquivo não encontrado no sistema de arquivos.', 'danger')
        return redirect(url_for('dashboard'))

@app.route('/delete_edital/<int:edital_id>')
@login_required
def delete_edital(edital_id):
    edital = Edital.query.get_or_404(edital_id)
    # Permite que o criador OU um admin delete
    if edital.creator_id != current_user.id and not current_user.is_admin():
        flash('Você não tem permissão para excluir este edital.', 'danger')
        return redirect(url_for('dashboard'))

    if edital.generated_filename: # Verifica se o nome do arquivo não é None ou vazio
        filepath = os.path.join(GENERATED_EDITALS_FOLDER, edital.generated_filename)
        if os.path.exists(filepath): # Verifica se o arquivo existe no disco
            try:
                os.remove(filepath) # Remove o arquivo físico
                flash('Arquivo associado excluído com sucesso!', 'info')
            except Exception as e:
                flash(f'Erro ao excluir o arquivo associado: {str(e)}', 'danger')
                app.logger.error(f"Erro ao excluir arquivo {filepath}: {e}", exc_info=True)
        else:
            flash('Aviso: Arquivo associado não encontrado no disco, mas o registro será removido.', 'warning')
    else:
        flash('Aviso: Nenhum arquivo associado registrado para este edital.', 'warning')

    try:
        db.session.delete(edital)
        db.session.commit()
        flash('Edital excluído com sucesso do banco de dados.', 'success')
    except Exception as e:
        flash(f'Erro ao excluir edital do banco de dados: {str(e)}', 'danger')
        db.session.rollback() # Em caso de erro, desfaz a transação
    return redirect(url_for('dashboard'))

# ================================================================
# NOVAS ROTAS DE ADMINISTRAÇÃO
# ================================================================

@app.route('/admin')
@admin_required
def admin_dashboard():
    total_users = User.query.count()
    total_editals = Edital.query.count()
    return render_template('admin_dashboard.html', total_users=total_users, total_editals=total_editals)

@app.route('/admin/editals')
@admin_required
def admin_all_editals():
    # Carrega TODOS os editais, ordenados pela data de criação
    all_editals = Edital.query.order_by(Edital.data_criacao.desc()).all()
    return render_template('admin_all_editals.html', editals=all_editals)

@app.route('/admin/users')
@admin_required
def admin_manage_users():
    users = User.query.all()
    return render_template('admin_manage_users.html', users=users)

@app.route('/admin/add_user', methods=['GET', 'POST'])
@admin_required
def admin_add_user():
    form = RegisterForm()
    if form.validate_on_submit():
        existing_user = User.query.filter_by(username=form.username.data).first()
        if existing_user:
            flash('Este nome de usuário já existe. Por favor, escolha outro.', 'danger')
        else:
            user = User(username=form.username.data, email=form.email.data)
            user.set_password(form.password.data)
            # Admin pode criar usuários comuns por aqui
            db.session.add(user)
            db.session.commit()
            flash(f'Usuário {form.username.data} adicionado com sucesso!', 'success')
            return redirect(url_for('admin_manage_users'))
    return render_template('admin_add_user.html', form=form)

@app.route('/admin/delete_user/<int:user_id>', methods=['POST'])
@admin_required
def admin_delete_user(user_id):
    user_to_delete = User.query.get_or_404(user_id)

    if user_to_delete.id == current_user.id:
        flash('Você não pode excluir sua própria conta de administrador.', 'danger')
        return redirect(url_for('admin_manage_users'))
    
    # Opcional: Impedir a exclusão do último admin
    if user_to_delete.is_admin() and User.query.filter_by(role='admin').count() == 1:
        flash('Não é possível excluir o último usuário administrador.', 'danger')
        return redirect(url_for('admin_manage_users'))

    try:
        # Excluir todos os editais criados por este usuário antes de excluir o usuário
        # Primeiro, exclua os arquivos DOCX associados aos editais do usuário
        editals_to_delete = Edital.query.filter_by(creator_id=user_to_delete.id).all()
        for edital_item in editals_to_delete:
            if edital_item.generated_filename:
                filepath = os.path.join(GENERATED_EDITALS_FOLDER, edital_item.generated_filename)
                if os.path.exists(filepath):
                    try:
                        os.remove(filepath)
                        print(f"[DEBUG] Arquivo {filepath} excluído ao remover usuário.")
                    except Exception as e:
                        print(f"[ERROR] Erro ao excluir arquivo {filepath} ao remover usuário: {e}")
        
        # Agora exclua os registros dos editais do banco de dados
        Edital.query.filter_by(creator_id=user_to_delete.id).delete()
        
        # Finalmente, exclua o usuário
        db.session.delete(user_to_delete)
        db.session.commit()
        flash(f'Usuário {user_to_delete.username} e seus editais excluídos com sucesso.', 'success')
    except Exception as e:
        flash(f'Erro ao excluir usuário: {str(e)}', 'danger')
        db.session.rollback()
    return redirect(url_for('admin_manage_users'))

# ================================================================
# 4. COMANDOS FLASK CLI
# ================================================================
@app.cli.command('init-db')
def init_db_command():
    """Cria as tabelas do banco de dados e um usuário admin inicial."""
    with app.app_context():
        db.create_all()
        # Cria um usuário admin se não existir
        if not User.query.filter_by(username='admin').first():
            admin_user = User(username='admin', email='admin@example.com', role='admin')
            admin_user.set_password('admin123') # Senha padrão para o admin
            db.session.add(admin_user)
            db.session.commit()
            click.echo("Usuário 'admin' criado com sucesso! Senha: admin123")
        else:
            click.echo("Usuário 'admin' já existe.")
        click.echo("Banco de dados inicializado.")

# ================================================================
# 5. FUNÇÃO PRINCIPAL - CONFIGURAÇÃO LOCAL E NUVEM
# ================================================================

import os

def create_default_admin():
    """Cria usuário admin padrão se não existir"""
    try:
        if not User.query.filter_by(username='admin').first():
            admin_user = User(username='admin', email='admin@edital.com', role='admin')
            admin_user.set_password('admin123')
            db.session.add(admin_user)
            db.session.commit()
            print("✅ Usuário admin criado: admin/admin123")
        else:
            print("ℹ️ Usuário admin já existe")
    except Exception as e:
        print(f"⚠️ Erro ao criar admin: {e}")

if __name__ == '__main__':
    # Cria as tabelas do banco de dados
    with app.app_context():
        db.create_all()
        create_default_admin()
    
    # Verifica se está rodando em produção (nuvem) ou desenvolvimento (local)
    if os.environ.get('RENDER'):
        # Configuração para RENDER (nuvem)
        print("🌐 RODANDO NA NUVEM (RENDER)")
        port = int(os.environ.get('PORT', 10000))
        app.run(host='0.0.0.0', port=port)
        
    elif os.environ.get('HEROKU'):
        # Configuração para HEROKU (nuvem)
        print("🌐 RODANDO NA NUVEM (HEROKU)")
        port = int(os.environ.get('PORT', 5000))
        app.run(host='0.0.0.0', port=port)
        
    else:
        # Configuração para DESENVOLVIMENTO LOCAL
        print("🏠 RODANDO LOCALMENTE")
        print("=" * 50)
        print("🔗 Acesse: http://127.0.0.1:5000")
        print("👤 Login: admin")
        print("🔑 Senha: admin123")
        print("=" * 50)
        print("💡 Para parar: Ctrl+C")
        print()
        

        app.run(host='127.0.0.1', port=5000, debug=True)
